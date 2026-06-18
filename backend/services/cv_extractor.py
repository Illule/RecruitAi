"""
services/cv_extractor.py
------------------------
Feature 2: Raw text extraction from uploaded CV files.

Supports:
  - PDF  → pdfplumber (handles multi-column, headers, footers)
  - DOCX → python-docx (paragraphs + tables)

Raises ValueError for unsupported file types.
"""

import io
import logging
from pathlib import Path

logger = logging.getLogger(__name__)


def extract_text(filename: str, file_bytes: bytes) -> str:
    """
    Extract plain text from a PDF or DOCX file.

    Args:
        filename:   Original filename (used to determine type).
        file_bytes: Raw file content as bytes.

    Returns:
        Extracted text as a single string (may be multi-page).

    Raises:
        ValueError: If the file type is not supported or extraction fails.
    """
    suffix = Path(filename).suffix.lower()

    if suffix == ".pdf":
        return _extract_pdf(filename, file_bytes)
    elif suffix in (".docx", ".doc"):
        return _extract_docx(filename, file_bytes)
    else:
        raise ValueError(
            f"Unsupported file type '{suffix}'. Only PDF and DOCX are accepted."
        )


# ─── Private helpers ──────────────────────────────────────────────────────────

def _extract_pdf(filename: str, file_bytes: bytes) -> str:
    """Extract text from a PDF using pdfplumber."""
    try:
        import pdfplumber  # type: ignore
    except ImportError as exc:
        raise RuntimeError("pdfplumber is not installed.") from exc

    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            pages = []
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    pages.append(text.strip())
            full_text = "\n\n".join(pages)

        if not full_text.strip():
            raise ValueError(f"PDF '{filename}' contained no extractable text.")

        logger.info("PDF '%s': extracted %d chars from %d pages",
                    filename, len(full_text), len(pdf.pages))
        return full_text

    except Exception as exc:
        logger.error("Failed to extract text from PDF '%s': %s", filename, exc)
        raise ValueError(f"Could not read PDF '{filename}': {exc}") from exc


def _extract_docx(filename: str, file_bytes: bytes) -> str:
    """Extract text from a DOCX using python-docx (paragraphs + tables)."""
    try:
        from docx import Document  # type: ignore
    except ImportError as exc:
        raise RuntimeError("python-docx is not installed.") from exc

    try:
        doc = Document(io.BytesIO(file_bytes))
        parts: list[str] = []

        # Paragraphs (main body + headings)
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)

        # Tables (often used for skills/education sections)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    parts.append(row_text)

        full_text = "\n".join(parts)

        if not full_text.strip():
            raise ValueError(f"DOCX '{filename}' contained no extractable text.")

        logger.info("DOCX '%s': extracted %d chars", filename, len(full_text))
        return full_text

    except Exception as exc:
        logger.error("Failed to extract text from DOCX '%s': %s", filename, exc)
        raise ValueError(f"Could not read DOCX '{filename}': {exc}") from exc
